"""
Unit tests for courtaccess/core/docx_translation.py

Functions under test:
  _split_fill_suffix   — pure: split label from trailing fill characters
  _is_translatable     — pure: decide whether a run text is worth translating
  _apply_translations  — write translations back into run objects
  translate_docx       — full pipeline (real .docx in tmp_path, mocked translator/reviewer)

Design notes:
  - _split_fill_suffix and _is_translatable are pure functions tested directly.
  - _apply_translations mutates python-docx Run objects; we use a real minimal
    .docx created in tmp_path rather than mocking the run XML interface.
  - translate_docx is tested with a real .docx and stub-mode Translator /
    LegalReviewer mocks. The output file is inspected via python-docx to
    verify runs were actually updated.
  - _collect_runs and _safe_set_run_text are exercised indirectly through
    translate_docx integration tests.
"""

from unittest.mock import MagicMock

import pytest
from docx import Document

from courtaccess.core.docx_translation import (
    _apply_translations,
    _is_translatable,
    _split_fill_suffix,
    translate_docx,
)

# ── Minimal DOCX factory ──────────────────────────────────────────────────────

def _make_docx(path, paragraphs=("Hello court.", "The defendant appeared.")):
    """Create a minimal .docx with the given paragraph texts."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(str(path))


def _make_docx_with_table(path):
    """Create a .docx with a 2x2 table containing translatable text."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Plaintiff name"
    table.cell(0, 1).text = "Defendant name"
    table.cell(1, 0).text = "Court date"
    table.cell(1, 1).text = "Case number"
    doc.save(str(path))


def _stub_translator():
    """MagicMock Translator that prefixes each text with [ES]."""
    t = MagicMock()
    t.batch_translate.side_effect = lambda texts, lang: [f"[ES] {x}" for x in texts]
    return t


def _stub_reviewer():
    """MagicMock LegalReviewer that passes translations through unchanged."""
    r = MagicMock()
    r.verify_batch.side_effect = lambda orig, trans, batch_size=16: trans
    return r


# ─────────────────────────────────────────────────────────────────────────────
# _split_fill_suffix — pure function
# ─────────────────────────────────────────────────────────────────────────────

class TestSplitFillSuffix:

    def test_no_fill_suffix_returns_text_and_empty(self):
        label, fill = _split_fill_suffix("Date:")
        assert label == "Date:"
        assert fill == ""

    def test_underscore_suffix_split_correctly(self):
        label, fill = _split_fill_suffix("Date:________________")
        assert label == "Date:"
        assert fill == "________________"

    def test_dash_suffix_split_correctly(self):
        label, fill = _split_fill_suffix("Name: -----------")
        assert label == "Name: "
        assert fill == "-----------"

    def test_trailing_space_included_in_fill(self):
        _, fill = _split_fill_suffix("Sign:______ ")
        assert fill.endswith(" ")

    def test_pure_fill_line_returns_empty_label(self):
        label, fill = _split_fill_suffix("____________")
        assert label == ""
        assert "____" in fill

    def test_three_underscores_not_detected(self):
        # _FILL_SUFFIX_RE requires {4,} — three chars not a fill suffix
        label, fill = _split_fill_suffix("A___")
        assert label == "A___"
        assert fill == ""

    def test_four_underscores_detected(self):
        label, fill = _split_fill_suffix("A____")
        assert label == "A"
        assert fill == "____"

    def test_plain_word_not_split(self):
        label, fill = _split_fill_suffix("Defendant")
        assert label == "Defendant"
        assert fill == ""


# ─────────────────────────────────────────────────────────────────────────────
# _is_translatable — pure function
# ─────────────────────────────────────────────────────────────────────────────

class TestIsTranslatable:

    def test_empty_string_not_translatable(self):
        assert _is_translatable("") is False

    def test_whitespace_only_not_translatable(self):
        assert _is_translatable("   ") is False

    def test_pure_fill_line_underscores_not_translatable(self):
        assert _is_translatable("__________") is False

    def test_pure_fill_line_dashes_not_translatable(self):
        assert _is_translatable("----------") is False

    def test_pure_numbers_not_translatable(self):
        assert _is_translatable("123 456") is False

    def test_number_with_punctuation_not_translatable(self):
        assert _is_translatable("(617) 555-1234") is False

    def test_url_not_translatable(self):
        assert _is_translatable("https://www.courts.ma.gov") is False

    def test_single_letter_not_translatable(self):
        # len(t.strip()) < 2 → False
        assert _is_translatable("A") is False

    def test_normal_sentence_is_translatable(self):
        assert _is_translatable("The defendant appeared.") is True

    def test_two_letter_word_is_translatable(self):
        assert _is_translatable("to") is True

    def test_mixed_text_with_numbers_is_translatable(self):
        assert _is_translatable("Case No. 1234") is True

    def test_legal_phrase_is_translatable(self):
        assert _is_translatable("Commonwealth of Massachusetts") is True


# ─────────────────────────────────────────────────────────────────────────────
# _apply_translations — mutates run objects
# ─────────────────────────────────────────────────────────────────────────────

class TestApplyTranslations:

    def test_raises_value_error_on_length_mismatch(self, tmp_path):
        docx_path = tmp_path / "doc.docx"
        _make_docx(docx_path, paragraphs=("Hello.", "World."))
        doc = Document(str(docx_path))
        runs = [r for p in doc.paragraphs for r in p.runs]
        with pytest.raises(ValueError, match="Mismatch"):
            _apply_translations(runs, ["one"])  # 2 runs, 1 translation

    def test_run_text_updated_in_place(self, tmp_path):
        docx_path = tmp_path / "doc.docx"
        _make_docx(docx_path, paragraphs=("Hello.",))
        doc = Document(str(docx_path))
        runs = [r for p in doc.paragraphs for r in p.runs]
        assert len(runs) == 1
        _apply_translations(runs, ["Hola."])
        assert runs[0].text == "Hola."

    def test_all_runs_updated(self, tmp_path):
        docx_path = tmp_path / "doc.docx"
        _make_docx(docx_path, paragraphs=("First.", "Second."))
        doc = Document(str(docx_path))
        runs = [r for p in doc.paragraphs for r in p.runs]
        _apply_translations(runs, ["Primero.", "Segundo."])
        assert runs[0].text == "Primero."
        assert runs[1].text == "Segundo."

    def test_empty_segments_and_translations_succeeds(self, tmp_path):
        # Zero runs, zero translations — should not raise
        _apply_translations([], [])


# ─────────────────────────────────────────────────────────────────────────────
# translate_docx — integration (real .docx, mocked translator/reviewer)
# ─────────────────────────────────────────────────────────────────────────────

class TestTranslateDocx:

    def test_output_file_created(self, tmp_path):
        src = tmp_path / "input.docx"
        out = tmp_path / "output.docx"
        _make_docx(src)
        translate_docx(str(src), str(out), _stub_translator(), _stub_reviewer(), "spa_Latn")
        assert out.exists()

    def test_returns_summary_dict_with_required_keys(self, tmp_path):
        src = tmp_path / "input.docx"
        out = tmp_path / "output.docx"
        _make_docx(src)
        result = translate_docx(str(src), str(out), _stub_translator(), _stub_reviewer(), "spa_Latn")
        assert "total_runs" in result
        assert "translated_runs" in result
        assert "fill_suffixes_found" in result

    def test_total_runs_matches_translatable_run_count(self, tmp_path):
        src = tmp_path / "input.docx"
        out = tmp_path / "output.docx"
        _make_docx(src, paragraphs=("Hello court.", "The defendant appeared."))
        result = translate_docx(str(src), str(out), _stub_translator(), _stub_reviewer(), "spa_Latn")
        assert result["total_runs"] == 2

    def test_translated_text_written_to_output_docx(self, tmp_path):
        src = tmp_path / "input.docx"
        out = tmp_path / "output.docx"
        _make_docx(src, paragraphs=("Hello court.",))
        translate_docx(str(src), str(out), _stub_translator(), _stub_reviewer(), "spa_Latn")
        result_doc = Document(str(out))
        texts = [r.text for p in result_doc.paragraphs for r in p.runs]
        assert any("[ES]" in t for t in texts)

    def test_empty_docx_returns_zero_runs(self, tmp_path):
        src = tmp_path / "empty.docx"
        out = tmp_path / "output.docx"
        _make_docx(src, paragraphs=())
        result = translate_docx(str(src), str(out), _stub_translator(), _stub_reviewer(), "spa_Latn")
        assert result["total_runs"] == 0
        assert result["translated_runs"] == 0

    def test_empty_docx_output_file_still_created(self, tmp_path):
        src = tmp_path / "empty.docx"
        out = tmp_path / "output.docx"
        _make_docx(src, paragraphs=())
        translate_docx(str(src), str(out), _stub_translator(), _stub_reviewer(), "spa_Latn")
        assert out.exists()

    def test_fill_suffix_reattached_after_translation(self, tmp_path):
        src = tmp_path / "fill.docx"
        out = tmp_path / "output.docx"
        # "Date:____________" — label translated, fill suffix reattached
        _make_docx(src, paragraphs=("Date:____________",))
        translate_docx(str(src), str(out), _stub_translator(), _stub_reviewer(), "spa_Latn")
        result_doc = Document(str(out))
        texts = [r.text for p in result_doc.paragraphs for r in p.runs]
        full_text = "".join(texts)
        assert "____" in full_text

    def test_fill_suffix_count_reported(self, tmp_path):
        src = tmp_path / "fill.docx"
        out = tmp_path / "output.docx"
        _make_docx(src, paragraphs=("Date:____________", "Sign:----------"))
        result = translate_docx(str(src), str(out), _stub_translator(), _stub_reviewer(), "spa_Latn")
        assert result["fill_suffixes_found"] == 2

    def test_translator_receives_clean_text_without_fill_suffix(self, tmp_path):
        src = tmp_path / "fill.docx"
        out = tmp_path / "output.docx"
        _make_docx(src, paragraphs=("Date:______",))
        mock_translator = _stub_translator()
        translate_docx(str(src), str(out), mock_translator, _stub_reviewer(), "spa_Latn")
        call_texts = mock_translator.batch_translate.call_args[0][0]
        # Fill suffix stripped before translation — only label sent
        assert call_texts == ["Date:"]

    def test_translator_called_with_correct_target_language(self, tmp_path):
        src = tmp_path / "input.docx"
        out = tmp_path / "output.docx"
        _make_docx(src)
        mock_translator = _stub_translator()
        translate_docx(str(src), str(out), mock_translator, _stub_reviewer(), "por_Latn")
        _, call_lang = mock_translator.batch_translate.call_args[0]
        assert call_lang == "por_Latn"

    def test_reviewer_verify_batch_called(self, tmp_path):
        src = tmp_path / "input.docx"
        out = tmp_path / "output.docx"
        _make_docx(src)
        mock_reviewer = _stub_reviewer()
        translate_docx(str(src), str(out), _stub_translator(), mock_reviewer, "spa_Latn")
        mock_reviewer.verify_batch.assert_called_once()

    def test_table_cells_translated(self, tmp_path):
        src = tmp_path / "table.docx"
        out = tmp_path / "output.docx"
        _make_docx_with_table(src)
        result = translate_docx(str(src), str(out), _stub_translator(), _stub_reviewer(), "spa_Latn")
        # 4 table cells, all translatable
        assert result["total_runs"] >= 4

    def test_non_translatable_runs_not_sent_to_translator(self, tmp_path):
        src = tmp_path / "input.docx"
        out = tmp_path / "output.docx"
        # Pure numbers and URL — neither should be translated
        _make_docx(src, paragraphs=("123456", "https://courts.ma.gov"))
        mock_translator = _stub_translator()
        result = translate_docx(str(src), str(out), mock_translator, _stub_reviewer(), "spa_Latn")
        assert result["total_runs"] == 0
        mock_translator.batch_translate.assert_not_called()
