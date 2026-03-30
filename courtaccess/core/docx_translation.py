"""
courtaccess/core/docx_translation.py

Translate a DOCX file in-place preserving all run-level formatting.

Preserves:
  - Run-level formatting (bold, italic, font size, color)
  - Paragraph & heading styles
  - Tables (cell by cell)
  - Images / embedded media (untouched — pass-through)
  - Fill lines  (_____ / -----) attached to labels
  - Soft line breaks (<w:br/>) inside runs
  - Content controls (<w:sdt>) such as fillable form fields

Reuses:
  - Translator.batch_translate()  from courtaccess.core.translation
  - LegalReviewer.verify_batch()  from courtaccess.core.legal_review

Source: Adapted from the DOCX Translator Colab notebook (Cell 10).
        Colab-specific code (files.upload, files.download, WORK_DIR)
        removed. Core logic preserved identically.
"""

from __future__ import annotations

import logging
import re

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph as DocxParagraph

from courtaccess.core.legal_review import LegalReviewer
from courtaccess.core.translation import Translator

logger = logging.getLogger(__name__)

# ── Helpers ───────────────────────────────────────────────────────────────────

# Trailing fill characters: "Date:___________" or "Name: ----------"
_FILL_SUFFIX_RE = re.compile(r"([_\-]{4,}\s*)$")


def _split_fill_suffix(text: str) -> tuple[str, str]:
    """
    Split a mixed label+fill run into (label, fill_suffix).

    e.g. 'Date:________________ ' -> ('Date:', '________________ ')
    Returns (text, '') if no fill suffix detected.

    Source: Cell 10 _split_fill_suffix() — identical logic.
    """
    m = _FILL_SUFFIX_RE.search(text)
    if m:
        return text[: m.start()], text[m.start() :]
    return text, ""


def _is_translatable(text: str) -> bool:
    """
    Determine if a run's text should be sent for translation.
    Skips blank strings, pure numbers, URLs, and pure fill lines.

    Source: Cell 10 _is_translatable() — identical logic.
    """
    t = text.strip()
    if not t:
        return False
    if re.fullmatch(r"[_\-]{2,}", t):  # pure fill line
        return False
    if re.fullmatch(r"[\d\s\.,\-\(\)/\\]+", t):  # pure numbers/punctuation
        return False
    if re.match(r"https?://", t):  # URLs
        return False
    return len(t) >= 2


def _collect_runs(doc: Document) -> tuple[list, list[str]]:
    """
    Walk ALL paragraphs including those nested inside <w:sdt>
    content controls (form fields) that doc.paragraphs misses.
    Uses .iter() to traverse the full XML tree.

    Returns:
        segments: list of python-docx Run objects (mutable references)
        texts:    list of the corresponding run text strings

    Source: Cell 10 _collect_runs() — identical logic.
    """
    segments: list = []
    texts: list[str] = []
    seen: set[int] = set()

    def _scan_para(para: DocxParagraph) -> None:
        eid = id(para._element)
        if eid in seen:
            return
        seen.add(eid)
        for run in para.runs:
            if _is_translatable(run.text):
                segments.append(run)
                texts.append(run.text)  # keep original spacing

    # Walk all <w:p> elements in the body — catches content controls
    for elem in doc.element.body.iter(qn("w:p")):
        _scan_para(DocxParagraph(elem, doc))

    # Walk tables separately (they have their own paragraph hierarchy)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _scan_para(para)

    return segments, texts


def _safe_set_run_text(run, new_text: str) -> None:
    """
    Write new_text into a run WITHOUT destroying <w:br/> elements.

    python-docx's run.text setter calls _r.clear_content() which wipes
    ALL children including <w:br/> soft-return markers, collapsing
    Shift+Enter line breaks. We update <w:t> directly instead.

    Source: Cell 10 _safe_set_run_text() — identical logic.
    """
    r = run._element
    t_elems = r.findall(qn("w:t"))
    if t_elems:
        t_elems[0].text = new_text
        # Preserve leading/trailing whitespace
        if new_text != new_text.strip():
            t_elems[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        # Remove extra <w:t> elements (rare but possible)
        for extra in t_elems[1:]:
            r.remove(extra)
    else:
        from lxml import etree

        t = etree.SubElement(r, qn("w:t"))
        t.text = new_text


def _apply_translations(segments: list, translated: list[str]) -> None:
    """
    Write translated strings back into run objects in-place.

    Source: Cell 10 _apply_translations() — identical logic.
    """
    if len(segments) != len(translated):
        raise ValueError(f"Mismatch: {len(segments)} runs vs {len(translated)} translations")
    for run, new_text in zip(segments, translated, strict=True):
        _safe_set_run_text(run, new_text)


# ── Public API ────────────────────────────────────────────────────────────────


def translate_docx(
    input_path: str,
    output_path: str,
    translator: Translator,
    reviewer: LegalReviewer,
    nllb_target: str,
) -> dict:
    """
    Translate a DOCX file preserving all run-level formatting.

    Steps (mirrors the Colab Cell 10 pipeline):
      1. Load the document with python-docx
      2. Collect all translatable runs (including inside content controls)
      3. Strip trailing fill suffixes before translation
      4. Translate via NLLB (Translator.batch_translate)
      5. Verify via Llama (LegalReviewer.verify_batch)
      6. Reattach fill suffixes
      7. Write verified translations back into the runs
      8. Save the document to output_path

    Args:
        input_path:   Path to the source .docx file
        output_path:  Path where the translated .docx will be saved
        translator:   Pre-loaded Translator instance (from courtaccess.core.translation)
        reviewer:     Pre-loaded LegalReviewer instance (from courtaccess.core.legal_review)
        nllb_target:  NLLB target language code, e.g. "spa_Latn" or "por_Latn"

    Returns:
        Summary dict:
          {
            "total_runs": int,
            "translated_runs": int,
            "fill_suffixes_found": int,
          }
    """
    # 1. Load
    doc = Document(input_path)

    # 2. Collect translatable runs
    segments, texts = _collect_runs(doc)
    total_runs = len(texts)
    logger.info("DOCX extracted %d translatable run(s) from %s", total_runs, input_path)

    if total_runs == 0:
        # Nothing to translate — save a copy and return
        doc.save(output_path)
        return {"total_runs": 0, "translated_runs": 0, "fill_suffixes_found": 0}

    # 3. Strip trailing fill suffixes before translation, reattach after
    clean_texts: list[str] = []
    fill_suffixes: list[str] = []
    for t in texts:
        clean, fill = _split_fill_suffix(t)
        clean_texts.append(clean)
        fill_suffixes.append(fill)

    fill_count = sum(1 for f in fill_suffixes if f)
    logger.info("Fill suffixes found: %d", fill_count)

    # 4. Translate (NLLB)
    logger.info("Running NLLB translation (%d segments) …", len(clean_texts))
    import time as _time

    t0 = _time.time()
    nllb_out = translator.batch_translate(clean_texts, nllb_target)
    nllb_changed = sum(1 for o, n in zip(clean_texts, nllb_out, strict=True) if o != n)
    logger.info("NLLB done — %d/%d segments changed", nllb_changed, len(clean_texts))

    # 5. Verify (Llama)
    logger.info("Running Llama legal verification …")
    verified = reviewer.verify_batch(clean_texts, nllb_out, batch_size=16)
    llama_corrections = sum(1 for n, v in zip(nllb_out, verified, strict=True) if n != v)
    elapsed = round(_time.time() - t0, 1)
    logger.info("Verification done — %d Llama correction(s) in %.1fs", llama_corrections, elapsed)

    # 6. Reattach fill suffixes
    verified = [t + fill for t, fill in zip(verified, fill_suffixes, strict=True)]

    # 7. Write back into the document runs
    logger.info("Reconstructing DOCX …")
    _apply_translations(segments, verified)

    # 8. Save
    doc.save(output_path)
    logger.info("Translated DOCX saved to %s", output_path)

    return {
        "total_runs": total_runs,
        "translated_runs": nllb_changed,
        "fill_suffixes_found": fill_count,
        "llama_corrections_count": llama_corrections,
        "elapsed_seconds": elapsed,
    }
